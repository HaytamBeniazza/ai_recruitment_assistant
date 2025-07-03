"""
File processor service for RecruitAI Pro
Handles PDF and DOCX resume text extraction with error handling and validation
"""

import os
import re
import mimetypes
import tempfile
import logging
from typing import Optional, Tuple, Dict, Any
from pathlib import Path

import PyPDF2
from docx import Document
from fastapi import HTTPException, UploadFile

# Configure logging
logger = logging.getLogger(__name__)

class FileProcessor:
    """
    Service for processing uploaded resume files
    Supports PDF and DOCX formats with robust error handling
    """
    
    # Supported file types
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc'}
    SUPPORTED_MIMETYPES = {
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword',
        'application/vnd.ms-word'
    }
    
    # File size limits (in bytes)
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    MIN_FILE_SIZE = 100  # 100 bytes
    
    # Text validation
    MIN_TEXT_LENGTH = 50  # Minimum characters for valid resume
    MAX_TEXT_LENGTH = 50000  # Maximum characters to prevent abuse
    
    def __init__(self, upload_directory: str = "uploads"):
        """Initialize file processor with upload directory"""
        self.upload_directory = Path(upload_directory)
        self.upload_directory.mkdir(exist_ok=True)
        
    async def process_resume_file(self, file: UploadFile) -> Dict[str, Any]:
        """
        Process uploaded resume file and extract text content
        
        Args:
            file: FastAPI UploadFile object
            
        Returns:
            Dict containing extracted text, metadata, and file info
            
        Raises:
            HTTPException: If file validation fails or processing errors occur
        """
        try:
            # Validate file
            self._validate_file(file)
            
            # Save file temporarily
            file_path = await self._save_uploaded_file(file)
            
            try:
                # Extract text based on file type
                extracted_text = self._extract_text_from_file(file_path, file.filename)
                
                # Validate extracted text
                self._validate_extracted_text(extracted_text)
                
                # Clean and process text
                cleaned_text = self._clean_text(extracted_text)
                
                # Extract metadata
                metadata = self._extract_metadata(cleaned_text, file.filename)
                
                return {
                    "success": True,
                    "filename": file.filename,
                    "file_size": file.size,
                    "file_path": str(file_path),
                    "extracted_text": cleaned_text,
                    "raw_text": extracted_text,
                    "text_length": len(cleaned_text),
                    "word_count": len(cleaned_text.split()),
                    "metadata": metadata,
                    "processing_notes": []
                }
                
            except Exception as e:
                # Clean up file on processing error
                if file_path.exists():
                    file_path.unlink()
                raise e
                
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Unexpected error processing file {file.filename}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Unexpected error processing file: {str(e)}"
            )
    
    def _validate_file(self, file: UploadFile) -> None:
        """Validate uploaded file meets requirements"""
        
        # Check file size
        if not file.size:
            raise HTTPException(status_code=400, detail="File is empty")
        
        if file.size < self.MIN_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File too small. Minimum size: {self.MIN_FILE_SIZE} bytes"
            )
        
        if file.size > self.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File too large. Maximum size: {self.MAX_FILE_SIZE // (1024*1024)}MB"
            )
        
        # Check file extension
        file_extension = Path(file.filename).suffix.lower()
        if file_extension not in self.SUPPORTED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type. Supported: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        # Check MIME type if available
        if file.content_type and file.content_type not in self.SUPPORTED_MIMETYPES:
            logger.warning(f"Unexpected MIME type: {file.content_type} for file {file.filename}")
    
    async def _save_uploaded_file(self, file: UploadFile) -> Path:
        """Save uploaded file to disk and return path"""
        
        # Generate safe filename
        safe_filename = self._generate_safe_filename(file.filename)
        file_path = self.upload_directory / safe_filename
        
        # Ensure unique filename
        counter = 1
        while file_path.exists():
            name_parts = Path(safe_filename).stem, counter, Path(safe_filename).suffix
            file_path = self.upload_directory / f"{name_parts[0]}_{name_parts[1]}{name_parts[2]}"
            counter += 1
        
        # Save file
        try:
            content = await file.read()
            with open(file_path, 'wb') as f:
                f.write(content)
            return file_path
        except Exception as e:
            logger.error(f"Error saving file {file.filename}: {str(e)}")
            raise HTTPException(status_code=500, detail="Error saving uploaded file")
    
    def _extract_text_from_file(self, file_path: Path, filename: str) -> str:
        """Extract text from file based on its type"""
        
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return self._extract_text_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._extract_text_from_docx(file_path)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {file_extension}"
                )
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {str(e)}")
            raise HTTPException(
                status_code=422,
                detail=f"Could not extract text from file. File may be corrupted or password-protected."
            )
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                raise ValueError("PDF is password-protected")
            
            # Extract text from all pages
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                    continue
        
        if not text_content:
            raise ValueError("No text content found in PDF")
        
        return '\n\n'.join(text_content)
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        text_content = []
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
        
        except Exception as e:
            raise ValueError(f"Error reading DOCX file: {str(e)}")
        
        if not text_content:
            raise ValueError("No text content found in document")
        
        return '\n'.join(text_content)
    
    def _validate_extracted_text(self, text: str) -> None:
        """Validate extracted text meets requirements"""
        
        if not text or not text.strip():
            raise HTTPException(
                status_code=422,
                detail="No text content found in file"
            )
        
        if len(text.strip()) < self.MIN_TEXT_LENGTH:
            raise HTTPException(
                status_code=422,
                detail=f"Resume too short. Minimum {self.MIN_TEXT_LENGTH} characters required"
            )
        
        if len(text) > self.MAX_TEXT_LENGTH:
            raise HTTPException(
                status_code=422,
                detail=f"Resume too long. Maximum {self.MAX_TEXT_LENGTH} characters allowed"
            )
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Remove non-printable characters except common ones
        text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
        
        # Normalize line breaks
        text = re.sub(r'\r\n?', '\n', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def _extract_metadata(self, text: str, filename: str) -> Dict[str, Any]:
        """Extract basic metadata from text content"""
        
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        
        # Basic statistics
        metadata = {
            "total_lines": len(lines),
            "total_words": len(text.split()),
            "total_characters": len(text),
            "filename": filename,
            "sections_detected": [],
            "potential_contact_info": self._extract_contact_info(text)
        }
        
        # Detect common resume sections
        section_keywords = {
            "summary": ["summary", "profile", "objective", "about"],
            "experience": ["experience", "employment", "work history", "career"],
            "education": ["education", "academic", "university", "degree"],
            "skills": ["skills", "technical", "competencies", "expertise"],
            "certifications": ["certifications", "certificates", "credentials"],
            "projects": ["projects", "portfolio", "achievements"]
        }
        
        text_lower = text.lower()
        for section, keywords in section_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                metadata["sections_detected"].append(section)
        
        return metadata
    
    def _extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract basic contact information patterns"""
        
        contact_info = {
            "email": None,
            "phone": None,
            "linkedin": None
        }
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact_info["email"] = email_match.group()
        
        # Phone pattern (various formats)
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact_info["phone"] = phone_match.group()
        
        # LinkedIn pattern
        linkedin_pattern = r'linkedin\.com/in/[a-zA-Z0-9-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact_info["linkedin"] = linkedin_match.group()
        
        return contact_info
    
    def _generate_safe_filename(self, filename: str) -> str:
        """Generate safe filename for storage"""
        
        # Keep only alphanumeric characters, dots, and hyphens
        safe_name = re.sub(r'[^a-zA-Z0-9.\-_]', '_', filename)
        
        # Ensure filename isn't too long
        if len(safe_name) > 100:
            name_part = safe_name[:80]
            extension = Path(safe_name).suffix
            safe_name = f"{name_part}{extension}"
        
        return safe_name
    
    def cleanup_file(self, file_path: str) -> bool:
        """Clean up uploaded file"""
        try:
            path = Path(file_path)
            if path.exists():
                path.unlink()
                return True
        except Exception as e:
            logger.error(f"Error cleaning up file {file_path}: {str(e)}")
        return False
    
    def get_file_info(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Get information about processed file"""
        try:
            path = Path(file_path)
            if not path.exists():
                return None
            
            stat = path.stat()
            return {
                "filename": path.name,
                "file_size": stat.st_size,
                "created_at": stat.st_ctime,
                "modified_at": stat.st_mtime,
                "file_extension": path.suffix.lower()
            }
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {str(e)}")
            return None

# Global instance for dependency injection
file_processor = FileProcessor() 